"""
Servicio de procesamiento de documentos.

Extrae texto de archivos PDF y Word (.docx) para su uso
en la generación de actas por IA.
"""

from __future__ import annotations

import io
from typing import Optional


def extract_text_from_pdf(content: bytes) -> str:
    """Extrae todo el texto de un archivo PDF."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise RuntimeError(
            "PyPDF2 no está instalado. Ejecuta: pip install PyPDF2"
        )

    reader = PdfReader(io.BytesIO(content))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def extract_text_from_docx(content: bytes) -> str:
    """Extrae todo el texto de un archivo Word (.docx)."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx no está instalado. Ejecuta: pip install python-docx"
        )

    doc = Document(io.BytesIO(content))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


def extract_text(filename: str, content: bytes) -> str:
    """Extrae texto según la extensión del archivo."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(content)
    if lower.endswith(".docx"):
        return extract_text_from_docx(content)
    if lower.endswith(".txt"):
        return content.decode("utf-8", errors="ignore")
    raise ValueError(
        f"Formato no soportado: {filename}. "
        "Se aceptan archivos PDF, Word (.docx) y texto plano (.txt)."
    )


def detect_format(filename: str) -> str:
    """Detecta el formato de origen del archivo."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith(".docx"):
        return "word"
    if lower.endswith(".txt"):
        return "transcripcion"
    return "manual"

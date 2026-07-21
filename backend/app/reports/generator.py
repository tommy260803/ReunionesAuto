"""Módulo de generación de reportes en PDF, Word y Excel."""

from datetime import datetime
from typing import Any

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False


def generate_report(
    report_type: str,
    format_type: str,
    data: dict[str, Any],
    title: str,
) -> bytes:
    """
    Genera un reporte en el formato solicitado.
    
    Args:
        report_type: Tipo de reporte (ANALISIS, EXPERIMENTO, EVALUACION)
        format_type: Formato (PDF, WORD, EXCEL)
        data: Datos del reporte
        title: Título del reporte
    
    Returns:
        Bytes del reporte generado
    """
    if format_type == "PDF":
        return generate_pdf_report(report_type, data, title)
    elif format_type == "WORD":
        return generate_word_report(report_type, data, title)
    elif format_type == "EXCEL":
        return generate_excel_report(report_type, data, title)
    else:
        raise ValueError(f"Formato no soportado: {format_type}")


def generate_pdf_report(
    report_type: str,
    data: dict[str, Any],
    title: str,
) -> bytes:
    """Genera un reporte en PDF."""
    if not REPORTLAB_AVAILABLE:
        raise ImportError("reportlab no está instalado. Instala con: pip install reportlab")
    
    from io import BytesIO
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Título
    title_style = styles["Heading1"]
    title_style.alignment = 1  # Center
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 12))
    
    # Fecha
    date_style = styles["Normal"]
    date_style.alignment = 1
    story.append(Paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", date_style))
    story.append(Spacer(1, 24))
    
    # Contenido según tipo
    if report_type == "ANALISIS":
        story.append(Paragraph("Resultados del Análisis", styles["Heading2"]))
        story.append(Spacer(1, 12))
        
        if "resultado" in data:
            story.append(Paragraph("Resultado Principal", styles["Heading3"]))
            story.append(Paragraph(str(data["resultado"]), styles["Normal"]))
            story.append(Spacer(1, 12))
        
        if "descriptivos" in data:
            story.append(Paragraph("Estadísticos Descriptivos", styles["Heading3"]))
            story.append(Paragraph(str(data["descriptivos"]), styles["Normal"]))
            story.append(Spacer(1, 12))
        
        if "interpretacion" in data:
            story.append(Paragraph("Interpretación", styles["Heading3"]))
            story.append(Paragraph(data["interpretacion"], styles["Normal"]))
    
    elif report_type == "EXPERIMENTO":
        story.append(Paragraph("Detalles del Experimento", styles["Heading2"]))
        story.append(Spacer(1, 12))
        
        for key, value in data.items():
            if key != "id":
                story.append(Paragraph(f"{key}: {value}", styles["Normal"]))
                story.append(Spacer(1, 6))
    
    elif report_type == "EVALUACION":
        story.append(Paragraph("Resultados de Evaluación", styles["Heading2"]))
        story.append(Spacer(1, 12))
        
        for key, value in data.items():
            if key != "id":
                story.append(Paragraph(f"{key}: {value}", styles["Normal"]))
                story.append(Spacer(1, 6))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_word_report(
    report_type: str,
    data: dict[str, Any],
    title: str,
) -> bytes:
    """Genera un reporte en Word."""
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError("python-docx no está instalado. Instala con: pip install python-docx")
    
    from io import BytesIO
    
    doc = Document()
    
    # Título
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fecha
    doc.add_paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    
    # Contenido según tipo
    if report_type == "ANALISIS":
        doc.add_heading("Resultados del Análisis", level=1)
        
        if "resultado" in data:
            doc.add_heading("Resultado Principal", level=2)
            doc.add_paragraph(str(data["resultado"]))
        
        if "descriptivos" in data:
            doc.add_heading("Estadísticos Descriptivos", level=2)
            doc.add_paragraph(str(data["descriptivos"]))
        
        if "interpretacion" in data:
            doc.add_heading("Interpretación", level=2)
            doc.add_paragraph(data["interpretacion"])
    
    elif report_type == "EXPERIMENTO":
        doc.add_heading("Detalles del Experimento", level=1)
        
        for key, value in data.items():
            if key != "id":
                doc.add_paragraph(f"{key}: {value}")
    
    elif report_type == "EVALUACION":
        doc.add_heading("Resultados de Evaluación", level=1)
        
        for key, value in data.items():
            if key != "id":
                doc.add_paragraph(f"{key}: {value}")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_excel_report(
    report_type: str,
    data: dict[str, Any],
    title: str,
) -> bytes:
    """Genera un reporte en Excel."""
    if not PANDAS_AVAILABLE:
        raise ImportError("pandas no está instalado. Instala con: pip install pandas openpyxl")
    
    from io import BytesIO
    
    # Convertir datos a DataFrame
    df_data = []
    
    if report_type == "ANALISIS":
        if "resultado" in data:
            df_data.append(["Resultado Principal", str(data["resultado"])])
        if "descriptivos" in data:
            df_data.append(["Estadísticos Descriptivos", str(data["descriptivos"])])
        if "interpretacion" in data:
            df_data.append(["Interpretación", data["interpretacion"]])
    else:
        for key, value in data.items():
            if key != "id":
                df_data.append([key, str(value)])
    
    df = pd.DataFrame(df_data, columns=["Campo", "Valor"])
    
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="Reporte", index=False)
    
    buffer.seek(0)
    return buffer.getvalue()


def check_format_availability(format_type: str) -> dict[str, Any]:
    """Verifica si el formato solicitado está disponible."""
    availability = {
        "PDF": REPORTLAB_AVAILABLE,
        "WORD": PYTHON_DOCX_AVAILABLE,
        "EXCEL": PANDAS_AVAILABLE,
    }
    
    return {
        "available": availability.get(format_type, False),
        "missing_dependency": None if availability.get(format_type, False) else get_missing_dependency(format_type),
    }


def get_missing_dependency(format_type: str) -> str | None:
    """Retorna la dependencia faltante para el formato."""
    dependencies = {
        "PDF": "reportlab",
        "WORD": "python-docx",
        "EXCEL": "pandas openpyxl",
    }
    return dependencies.get(format_type)
